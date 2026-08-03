"""
报告导出工具 - 支持 PDF 和 Word 格式
"""
import io
import logging
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ExportClient:
    """报告导出客户端"""

    def __init__(self):
        pass

    def export_to_pdf(self, content_md: str, title: str = "潜在需求分析报告") -> bytes:
        """
        将 Markdown 内容转换为 PDF

        参数：
            content_md: Markdown 格式的报告内容
            title: 报告标题

        返回：
            bytes: PDF 文件的二进制数据
        """
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
        except OSError as e:
            raise RuntimeError(
                "PDF 导出需要 Pango/Cairo 系统库。"
                "Windows 上请安装 GTK3 运行时：https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer"
            ) from e

        # 将 Markdown 转换为 HTML
        html_content = self._markdown_to_html(content_md, title)

        # 配置中文字体
        font_config = FontConfiguration()

        # 创建 CSS 样式
        # 注意：不使用 Google Fonts @import（国内网络不可达，会导致 weasyprint 挂起）
        # 使用系统预装的中文字体（Dockerfile 已安装 fonts-wqy-zenhei 文泉驿正黑）
        css = CSS(string="""
            body {
                font-family: 'WenQuanYi Zen Hei', 'SimSun', 'Microsoft YaHei', sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }

            h1, h2, h3, h4, h5, h6 {
                font-family: 'WenQuanYi Zen Hei', 'SimHei', 'Microsoft YaHei', sans-serif;
                color: #1a1a1a;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
            }

            h1 {
                font-size: 18pt;
                border-bottom: 2px solid #2563eb;
                padding-bottom: 0.3em;
            }

            h2 {
                font-size: 14pt;
                border-bottom: 1px solid #e5e7eb;
                padding-bottom: 0.3em;
            }

            h3 {
                font-size: 12pt;
            }

            p {
                margin: 0.8em 0;
            }

            ul, ol {
                margin: 0.8em 0;
                padding-left: 2em;
            }

            li {
                margin: 0.3em 0;
            }

            strong {
                font-weight: 600;
            }

            code {
                background-color: #f3f4f6;
                padding: 0.2em 0.4em;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 0.9em;
            }

            blockquote {
                border-left: 4px solid #e5e7eb;
                padding-left: 1em;
                margin-left: 0;
                color: #6b7280;
                font-style: italic;
            }

            table {
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
            }

            th, td {
                border: 1px solid #e5e7eb;
                padding: 0.5em 1em;
                text-align: left;
            }

            th {
                background-color: #f3f4f6;
                font-weight: 600;
            }

            a {
                color: #2563eb;
                text-decoration: none;
            }

            @page {
                size: A4;
                margin: 2.5cm 2cm;
            }

            @page:first {
                margin-top: 3cm;
            }
        """, font_config=font_config)

        # 生成 PDF（pdf/a-3u 变体强制输出 ToUnicode CMap，保证中文文本层可复制/搜索）
        pdf_buffer = io.BytesIO()
        html = HTML(string=html_content)
        html.write_pdf(pdf_buffer, css=css, pdf_variant="pdf/a-3u")
        pdf_buffer.seek(0)

        return pdf_buffer.getvalue()

    def export_to_word(self, content_md: str, title: str = "潜在需求分析报告") -> bytes:
        """
        将 Markdown 内容转换为 Word 文档

        参数：
            content_md: Markdown 格式的报告内容
            title: 报告标题

        返回：
            bytes: Word 文档的二进制数据
        """
        doc = Document()

        # 设置默认样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 解析并添加 Markdown 内容
        self._add_markdown_to_doc(doc, content_md)

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _markdown_to_html(self, content_md: str, title: str) -> str:
        """简单的 Markdown 转 HTML 转换"""
        # 转义 HTML 特殊字符
        html = content_md.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        # 处理标题
        html = self._process_headers(html)

        # 处理粗体
        html = html.replace("**", "<strong>").replace("**", "</strong>")

        # 处理列表
        html = self._process_lists(html)

        # 处理换行
        html = html.replace("\n\n", "</p><p>")

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
</head>
<body>
{html}
</body>
</html>"""

    def _process_headers(self, html: str) -> str:
        """处理 Markdown 标题"""
        lines = html.split("\n")
        result = []
        for line in lines:
            if line.startswith("### "):
                result.append(f"<h3>{line[4:]}</h3>")
            elif line.startswith("## "):
                result.append(f"<h2>{line[3:]}</h2>")
            elif line.startswith("# "):
                result.append(f"<h1>{line[2:]}</h1>")
            else:
                result.append(line)
        return "\n".join(result)

    def _process_lists(self, html: str) -> str:
        """处理 Markdown 列表"""
        lines = html.split("\n")
        result = []
        in_list = False

        for line in lines:
            if line.strip().startswith("- "):
                if not in_list:
                    result.append("<ul>")
                    in_list = True
                content = line.strip()[2:]
                result.append(f"<li>{content}</li>")
            else:
                if in_list:
                    result.append("</ul>")
                    in_list = False
                result.append(line)

        if in_list:
            result.append("</ul>")

        return "\n".join(result)

    def _add_markdown_to_doc(self, doc: Document, content_md: str):
        """将 Markdown 内容添加到 Word 文档"""
        lines = content_md.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # 处理标题
            if line.startswith("### "):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.italic = True
            else:
                # 处理内联格式
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

    def _add_formatted_text(self, paragraph, text: str):
        """添加带格式的文本到段落"""
        # 简单处理粗体
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:  # 粗体部分
                run = paragraph.add_run(part)
                run.bold = True
            else:
                paragraph.add_run(part)
